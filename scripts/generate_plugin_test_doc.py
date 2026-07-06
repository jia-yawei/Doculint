from pathlib import Path
from io import BytesIO

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
OUTPUT_DOC = OUTPUT_DIR / "需求规格说明-插件全功能测试文档.docx"


def set_run_font(run, name="Arial", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor(0, 0, 0)

    if "Caption" in doc.styles:
        caption = doc.styles["Caption"]
        caption.font.name = "Arial"
        caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        caption.font.size = Pt(11)

    if "插件测试提示" not in [s.name for s in doc.styles]:
        custom = doc.styles.add_style("插件测试提示", WD_STYLE_TYPE.PARAGRAPH)
        custom.base_style = doc.styles["Normal"]
        custom.font.name = "Arial"
        custom._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        custom._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        custom.font.size = Pt(10.5)
        custom.font.color.rgb = RGBColor(80, 80, 80)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("需求规格说明-插件全功能测试文档")
    set_run_font(run, size=22, bold=False)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(18)
    run2 = p2.add_run("用于验证“搞快点”选项卡中的标题、表格、图片题注、批量功能、需求追踪、抓取与注入等能力")
    set_run_font(run2, size=11, color=(90, 90, 90))


def add_meta_table(doc):
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Cm(3.5), Cm(11.5)]
    labels = [
        ("文档类型", "需求规格说明"),
        ("测试用途", "插件回归验证、界面联调、题注与表格处理验证"),
        ("推荐文档组", "插件测试组"),
        ("批量替换关键字", "项目的相关组织机构 / 系统功能 / 软件功能 / 测试关键字A / 测试关键字B"),
        ("内容抓取建议", "可在本表、1.2.1、图片测试区、长表格测试区中任意选中片段进行抓取"),
    ]
    for row_idx, (left, right) in enumerate(labels):
        for col_idx, text in enumerate((left, right)):
            cell = table.cell(row_idx, col_idx)
            cell.width = widths[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_run_font(run, size=10.5, bold=(col_idx == 0))


def add_toc_field(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "目录将在 Word 中更新后显示"
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    tip = doc.add_paragraph("测试提示：打开文档后可先执行“刷新目录”。", style="插件测试提示")
    tip.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading(doc, text, level):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, size=11)


def build_test_image(title, subtitle, color):
    image = Image.new("RGB", (1280, 720), color)
    draw = ImageDraw.Draw(image)
    font_big = ImageFont.load_default()
    font_small = ImageFont.load_default()
    draw.rounded_rectangle((60, 60, 1220, 660), radius=36, outline=(255, 255, 255), width=4)
    draw.text((120, 150), title, fill=(255, 255, 255), font=font_big)
    draw.text((120, 220), subtitle, fill=(240, 240, 240), font=font_small)
    draw.rectangle((120, 310, 580, 520), outline=(255, 255, 255), width=4)
    draw.ellipse((720, 300, 1080, 560), outline=(255, 255, 255), width=4)
    draw.text((150, 370), "图形区域A", fill=(255, 255, 255), font=font_small)
    draw.text((820, 405), "图形区域B", fill=(255, 255, 255), font=font_small)
    stream = BytesIO()
    image.save(stream, format="PNG")
    stream.seek(0)
    return stream


def add_picture_zone(doc):
    add_heading(doc, "2 图片与题注测试区", 1)
    add_body(doc, "下面两张图用于验证“插入图片题注”“更新图片题注”“题注引用”“窗格显示”等功能。")

    for idx, (title, subtitle, color) in enumerate(
        [
            ("插件测试示意图 A", "用于测试图题注、引用、抓取和居中样式。", (61, 108, 179)),
            ("插件测试示意图 B", "用于测试后续更新题注和图表导航。", (49, 133, 91)),
        ],
        start=1,
    ):
        stream = build_test_image(title, subtitle, color)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(stream, width=Inches(5.8))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(f"此处用于插入图片题注 {idx}")
        set_run_font(cap_run, size=11, color=(90, 90, 90))


def add_table_caption_placeholder(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, size=11, color=(90, 90, 90))


def set_cell(cell, text, bold=False, center=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)


def add_small_tables(doc):
    add_heading(doc, "3 表格测试区", 1)
    add_body(doc, "本节用于验证表格题注、表格选取、自动序号、跨页表格、续表拆分与合并。")

    add_heading(doc, "3.1 普通信息表", 2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    data = [
        ["项目", "内容", "说明"],
        ["测试关键字A", "需要在批量替换中命中", "用于查找与替换验证"],
        ["测试关键字B", "需要在批量替换中命中", "用于活动组多文件验证"],
        ["项目的相关组织机构", "本行用于章节关键字搜索", "不要做全字匹配"],
    ]
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            set_cell(table.cell(r_idx, c_idx), value, bold=(r_idx == 0), center=(c_idx == 0 and r_idx == 0))
    add_table_caption_placeholder(doc, "此处用于插入表格题注 1")

    add_heading(doc, "3.2 自动序号测试表", 2)
    add_body(doc, "将光标放在第一列任意单元格，可验证快速工具中的自动序号是否按纵向列表插入纯数字编号。")
    table2 = doc.add_table(rows=7, cols=3)
    table2.style = "Table Grid"
    seq_data = [
        ["序号", "检查项", "预期结果"],
        ["", "验证文档组下拉显示活动组", "下拉中能看到当前活动组名称"],
        ["", "验证内容抓取", "选中内容后可输入标题并保存到组中"],
        ["", "验证内容注入", "双击抓取标题后可在光标处注入原格式内容"],
        ["", "验证图片题注", "可在当前段落插入图x并后续更新"],
        ["", "验证表格题注", "可在当前段落插入表x并后续更新"],
        ["", "验证标识窗格", "仅受支持文档类型显示正常内容"],
    ]
    for r_idx, row in enumerate(seq_data):
        for c_idx, value in enumerate(row):
            set_cell(table2.cell(r_idx, c_idx), value, bold=(r_idx == 0), center=(c_idx == 0))


def add_long_table(doc):
    add_heading(doc, "3.3 跨页长表格测试", 2)
    add_body(doc, "下表包含较多行，用于测试“下个跨页表格”“按续表拆分”“合并续表”“设置为表头”等功能。")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["序号", "需求编号", "需求名称", "验证说明", "备注"]
    for idx, header in enumerate(headers):
        set_cell(table.cell(0, idx), header, bold=True, center=True)

    for i in range(1, 38):
        row = table.add_row().cells
        set_cell(row[0], str(i), center=True)
        set_cell(row[1], f"SRS-FUNC-{i:03d}", center=True)
        set_cell(row[2], f"功能需求条目 {i}")
        set_cell(
            row[3],
            "用于验证跨页长表格处理、表头重复、续表拆分以及合并逻辑。" if i % 2 else "用于验证表格选择、导航和样式一致性。",
        )
        set_cell(row[4], "奇数行" if i % 2 else "偶数行", center=True)
    add_table_caption_placeholder(doc, "此处用于插入表格题注 2")


def add_requirement_sections(doc):
    add_heading(doc, "1 说明", 1)
    add_body(doc, "本测试文档模拟需求规格说明书，包含用于需求追踪、章节定位、题注、表格和批量处理的测试内容。")

    add_heading(doc, "1.1 文档目的", 2)
    add_body(doc, "用于在“搞快点”选项卡下对常用功能进行一次完整回归验证，尤其是章节定位、文档组、题注、表格处理和抓取注入相关功能。")

    add_heading(doc, "1.2 章节定位测试区", 2)
    add_body(doc, "下面三个三级标题故意使用常见目标章节名称，用于验证章节搜索、定位高亮、内容提取与替换。")

    add_heading(doc, "1.2.1 系统功能", 3)
    add_body(doc, "系统功能测试内容第一段。这里包含测试关键字A、测试关键字B，以及批量替换前后的比对文本。")
    add_body(doc, "系统功能测试内容第二段。建议在这里测试内容抓取、标题样式、章节定位和段落复制。")
    add_bullets(
        doc,
        [
            "系统功能点 1：支持文档组、活动组和多文档管理。",
            "系统功能点 2：支持目录刷新、题注更新与图表导航。",
            "系统功能点 3：支持内容抓取、注入与组内共享。",
        ],
    )

    add_heading(doc, "1.2.2 软件功能", 3)
    add_body(doc, "软件功能测试内容第一段。该段用于验证从章节标题中搜索并提取正文内容到内存的逻辑。")
    add_body(doc, "软件功能测试内容第二段。定位后可直接覆盖目标文档对应章节内容，并保持原有段落格式。")
    add_numbered(
        doc,
        [
            "打开目标文档并定位到同名章节。",
            "删除原章节正文内容，不删除章节标题。",
            "将源章节内容原格式注入目标章节中。",
        ],
    )

    add_heading(doc, "1.2.3 项目的相关组织机构", 3)
    add_body(doc, "项目的相关组织机构测试内容第一段。注意本标题可能包含“项目的相关机构”的前缀关系，因此不能做全字匹配。")
    add_body(doc, "项目的相关组织机构测试内容第二段。这里可以验证取消全字匹配后是否还能从章节标题快速定位。")
    add_bullets(
        doc,
        [
            "组织机构 A：总体组，负责章节审查。",
            "组织机构 B：开发组，负责功能实现。",
            "组织机构 C：测试组，负责插件回归验证。",
        ],
    )

    add_heading(doc, "1.3 需求追踪测试提示", 2)
    add_body(doc, "如果用本文件测试需求追踪，建议重点检查最后两条需求名称的提取，确认不会误取表格第三列内容。")
    trace_table = doc.add_table(rows=6, cols=4)
    trace_table.style = "Table Grid"
    trace_rows = [
        ["序号", "需求名称", "需求描述", "来源"],
        ["1", "文档组持久化", "文档组信息关闭后仍可再次加载", "3.1 普通信息表"],
        ["2", "内容抓取与注入", "抓取内容保留原格式，可跨组内文档复用", "1.2.1 系统功能"],
        ["3", "图片题注更新", "更新时仅刷新已有题注域，不插入或删除其他内容", "2 图片与题注测试区"],
        ["4", "表格题注更新", "逻辑与图片题注一致，仅刷新现有题注域", "3 表格测试区"],
        ["5", "最后一条需求名称在第二列", "用于验证最后两条名称不应从第三列提取", "本表第二列"],
    ]
    for r_idx, row in enumerate(trace_rows):
        for c_idx, value in enumerate(row):
            set_cell(trace_table.cell(r_idx, c_idx), value, bold=(r_idx == 0), center=(c_idx == 0))


def add_capture_zone(doc):
    add_heading(doc, "4 内容抓取与注入测试区", 1)
    add_body(doc, "本节中的段落、表格和图片都可以用于测试“内容抓取”与“内容注入”。建议先将文档加入活动组并保存。")
    add_body(doc, "抓取建议一：选中本段完整内容，命名为“抓取示例-正文段落”。")
    add_body(doc, "抓取建议二：选中下方两列表格，命名为“抓取示例-检查清单表格”。")

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    capture_rows = [
        ["检查项", "操作说明"],
        ["抓取", "选中一段或一块原格式内容后点击“内容抓取”。"],
        ["管理", "在“抓取管理”中查看标题、修改标题、预览内容。"],
        ["注入", "点击“内容注入”后双击标题，将内容按原格式插入当前光标处。"],
        ["组限制", "只有打开同一文档组中的文档时，才能看到并注入抓取内容。"],
    ]
    for r_idx, row in enumerate(capture_rows):
        for c_idx, value in enumerate(row):
            set_cell(table.cell(r_idx, c_idx), value, bold=(r_idx == 0))


def add_navigation_zone(doc):
    add_heading(doc, "5 导航与显示测试区", 1)
    add_body(doc, "本节用于测试目录刷新、图表导航、总页数、标识窗格以及快速定位。")
    add_bullets(
        doc,
        [
            "因为当前文档名称为“需求规格说明-插件全功能测试文档”，应属于受支持的标识窗格文档类型。",
            "可以先尝试刷新目录，再测试图表项的导航和题注刷新。",
            "如果复制本文件并改名为其他类型，再打开标识窗格时应显示“不支持此文档”。",
        ],
    )

    add_heading(doc, "5.1 补充正文", 2)
    for idx in range(1, 10):
        add_body(
            doc,
            f"补充测试段落 {idx}：用于拉长文档页数，便于验证页码、导航、标题定位、抓取与注入后的页面表现是否稳定。",
        )


def add_section_break(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title(doc)
    add_meta_table(doc)
    add_section_break(doc)
    add_heading(doc, "目录", 1)
    add_toc_field(doc)
    add_requirement_sections(doc)
    add_picture_zone(doc)
    add_small_tables(doc)
    add_long_table(doc)
    add_capture_zone(doc)
    add_navigation_zone(doc)
    doc.save(OUTPUT_DOC)
    return OUTPUT_DOC


if __name__ == "__main__":
    output = build_document()
    print(output)
